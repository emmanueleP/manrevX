from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from .images_manager import images_manager
from .settings import manrev_settings
import os
from PIL import Image
import io

class DocumentLayout:
    def __init__(self, document):
        self.document = document
        self.sections = document.sections
        self.section = self.sections[0]
    
    def set_margins(self):
        """Imposta i margini del documento"""
        self.section.top_margin = Cm(2)
        self.section.bottom_margin = Cm(2)
        self.section.left_margin = Cm(2)
        self.section.right_margin = Cm(2)

    def add_header(self, title, year, number):
        """Aggiunge l'intestazione del documento"""
        # Aggiungi immagine sede se presente
        sede_path = manrev_settings.current_settings.get("sede_image", "")
        if sede_path and os.path.exists(sede_path):
            sede_para = self.document.add_paragraph()
            sede_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                sede_para.add_run().add_picture(sede_path, width=Cm(6))
                sede_para.add_run("\n")
            except Exception as e:
                print(f"Errore nel caricamento dell'immagine della sede: {e}")
        
        # Titolo e numero
        header = self.document.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = header.add_run(f"{title} N. {number}/{str(manrev_settings.current_settings.get('year'))}")
        title_run.bold = True
        title_run.font.size = Pt(14)

    def add_details_table(self, details):
        """Aggiunge la tabella dei dettagli"""
        table = self.document.add_table(rows=len(details), cols=2)
        table.style = 'Table Grid'
        
        for i, (key, value) in enumerate(details.items()):
            # Cella sinistra (chiave)
            cell = table.cell(i, 0)
            if key == 'Importo':
                cell.text = 'Importo in €'
            else:
                cell.text = key
            cell.paragraphs[0].runs[0].bold = True
            
            # Cella destra (valore)
            cell = table.cell(i, 1)
            if key == 'Importo':
                cell.text = f"{value} €"
            else:
                cell.text = str(value)

    def add_amount_text(self, amount_text):
        """Aggiunge l'importo in lettere"""
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = para.add_run(f"Importo in lettere: {amount_text}")
        run.bold = True
        run.font.size = Pt(11)

    def add_signatures(self, signatures):
        """Aggiunge la sezione firme con immagini"""
        # Aggiungi spazio prima delle firme
        self.document.add_paragraph()
        
        # Crea tabella per le firme (3 colonne)
        table = self.document.add_table(rows=2, cols=3)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Dizionario che mappa i ruoli ai file delle firme
        signature_files = {
            'Il Tesoriere': manrev_settings.current_settings.get('firme', {}).get('tesoriere_firma', ''),
            'Il Presidente': manrev_settings.current_settings.get('firme', {}).get('presidente_firma', ''),
            "L'Addetto Contabile": manrev_settings.current_settings.get('firme', {}).get('addetto_firma', '')
        }
        
        # Prima riga: Immagini delle firme
        for idx, (role, name) in enumerate(signatures.items()):
            cell = table.cell(0, idx)
            paragraph = cell.paragraphs[0]
            
            # Se esiste un'immagine della firma, aggiungila
            signature_path = signature_files.get(role, '')
            if signature_path and os.path.exists(signature_path):
                try:
                    # Ridimensiona l'immagine
                    with Image.open(signature_path) as img:
                        # Converti in RGB se necessario
                        if img.mode != 'RGB':
                            img = img.convert('RGB')
                        # Ridimensiona mantenendo l'aspect ratio
                        max_width = 100  # pixel
                        ratio = max_width / img.width
                        new_size = (max_width, int(img.height * ratio))
                        img = img.resize(new_size, Image.Resampling.LANCZOS)
                        
                        # Salva in un buffer
                        img_byte_arr = io.BytesIO()
                        img.save(img_byte_arr, format='PNG')
                        img_byte_arr.seek(0)
                        
                        # Aggiungi al documento
                        run = paragraph.add_run()
                        run.add_picture(img_byte_arr, width=Inches(1))
                except Exception as e:
                    print(f"Errore nel caricare la firma {role}: {e}")
            
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Seconda riga: Nomi
        for idx, (role, name) in enumerate(signatures.items()):
            cell = table.cell(1, idx)
            paragraph = cell.paragraphs[0]
            paragraph.add_run(f"{role}\n{name}")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Formattazione tabella
        table.allow_autofit = True
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Aggiungi spazio dopo la tabella
        self.document.add_paragraph()

    def add_footer(self, place, date):
        """Aggiunge il piè di pagina"""
        footer = self.document.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
        footer.add_run(f"\n{place}, {date}") 